import re
import os
import locale
import shutil
import logging
import tkinter as tk
import win32com.client
from docx import Document
from docx.shared import Pt
from datetime import datetime
from num2words import num2words
from openpyxl import load_workbook
from tkinter import messagebox, ttk
from collections import defaultdict
from cores import DARK_GREEN , DARK_RED ,RESET
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def formatar_data(data_str):
    data = datetime.strptime(data_str, "%Y-%m-%d %H:%M:%S")
    return data.strftime("%d/%m/%Y")
def substituir_palavras_documento(doc_path, substituicoes, output_path):
    documento = Document(doc_path)
    
    for paragrafo in documento.paragraphs:
        for run in paragrafo.runs:
            for palavra_antiga, palavra_nova in substituicoes.items():
                if palavra_antiga in run.text:
                    print(f"{DARK_GREEN}Substituindo '{palavra_antiga}' por \n'{palavra_nova}'\n {RESET}")
                    run.text = run.text.replace(palavra_antiga, palavra_nova)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        for palavra_antiga, palavra_nova in substituicoes.items():
                            if palavra_antiga in run.text:
                                print(f"{DARK_GREEN}Substituindo '{palavra_antiga}' por \n'{palavra_nova}'\n {RESET}")
                                run.text = run.text.replace(palavra_antiga, palavra_nova)

    documento.save(output_path)
    print(f"{DARK_GREEN}Substituição concluída. Documento salvo em {output_path}{RESET}")
def gerar_data_atual():
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    
    data = datetime.now()
    dia = data.day
    mes = meses[data.month - 1]
    ano = data.year
    
    return f"Palmas - TO, {dia} de {mes} de {ano}."
def ler_celula_excel(celula):
    caminho_arquivo = 'integracao.xlsx'
    nome_planilha = 'dados' 
    # Abre o arquivo Excel com data_only=True para ler o valor das fórmulas
    workbook = load_workbook(caminho_arquivo, data_only=True)
    # Seleciona a planilha especificada
    planilha = workbook[nome_planilha]
    # Obtém o valor da célula especificada
    valor = planilha[celula].value
    return valor
def formatar_valor(valor):
    valor = float(valor)
    valor_formatado = "R$ {:,.2f}".format(valor)
    valor_formatado = valor_formatado.replace(",", "X").replace(".", ",").replace("X", ".")
    return valor_formatado
def abrir_arquivo_word(caminho_arquivo):
    try:
        # Abre o arquivo Word no aplicativo padrão do sistema
        os.startfile(caminho_arquivo)
    except Exception as e:
        print(f"{DARK_RED}Erro ao abrir o arquivo: {e}{RESET}")
def valor_por_extenso(valor_str):
    # Remover o 'R$' e os pontos, substituir a vírgula por ponto para converter em float
    valor_numerico = float(re.sub(r'[^0-9,]', '', valor_str).replace(',', '.'))
    
    # Parte inteira e parte decimal
    parte_inteira = int(valor_numerico)
    parte_decimal = int(round((valor_numerico - parte_inteira) * 100))

    # Converter parte inteira e decimal para extenso
    extenso_inteira = num2words(parte_inteira, lang='pt_BR')
    extenso_decimal = num2words(parte_decimal, lang='pt_BR')
    
    # Montar a resposta por extenso
    if parte_decimal > 0:
        valor_extenso = f"{extenso_inteira} reais e {extenso_decimal} centavos"
    else:
        valor_extenso = f"{extenso_inteira} reais"

    return valor_extenso
def substituir_ponto_por_virgulas(varivel,casas_decimais):
    
    # Formata o valor com a quantidade desejada de casas decimais
    formatted_value = f"{varivel:.{casas_decimais}f}"
    
    # Substitui pontos por vírgulas
    formatted_value = formatted_value.replace('.', ',')

    return formatted_value
def escolher_e_ler_arquivo_txt(excel_path):
    print(f"{DARK_GREEN}[DEBUG] Caminho do arquivo Excel recebido: {excel_path}{RESET}")
    current_dir = os.path.dirname(excel_path)
    print(f"{DARK_GREEN}[DEBUG] Diretório inicial: {current_dir}{RESET}")

    all_contents = []

    # Verifica o diretório atual
    print(f"{DARK_GREEN}[DEBUG] Verificando diretório: {current_dir}{RESET}")
    for file in os.listdir(current_dir):
        print(f"{DARK_GREEN}[DEBUG] Arquivo encontrado: {file}{RESET}")
        if file.endswith(".txt"):
            found_path = os.path.join(current_dir, file)
            print(f"{DARK_GREEN}[DEBUG] Arquivo .txt encontrado: {found_path}{RESET}")
            with open(found_path, 'r', encoding='utf-8') as txt_file:
                content = txt_file.read()
            print(f"{DARK_GREEN}[DEBUG] Conteúdo do arquivo .txt lido com sucesso.{RESET}")
            all_contents.append(content)

    # Sobe um diretório e verifica novamente
    parent_dir = os.path.abspath(os.path.join(current_dir, os.pardir))
    print(f"{DARK_GREEN}[DEBUG] Subindo para o diretório pai: {parent_dir}{RESET}")
    if parent_dir != current_dir:  # Verifica se pode subir
        print(f"{DARK_GREEN}[DEBUG] Verificando diretório pai: {parent_dir}{RESET}")
        for file in os.listdir(parent_dir):
            print(f"{DARK_GREEN}[DEBUG] Arquivo encontrado: {file}{RESET}")
            if file.endswith(".txt"):
                found_path = os.path.join(parent_dir, file)
                print(f"{DARK_GREEN}[DEBUG] Arquivo .txt encontrado: {found_path}{RESET}")
                with open(found_path, 'r', encoding='utf-8') as txt_file:
                    content = txt_file.read()
                print(f"{DARK_GREEN}[DEBUG] Conteúdo do arquivo .txt lido com sucesso.{RESET}")
                all_contents.append(content)

    if all_contents:
        print(f"{DARK_GREEN}[DEBUG] Todos os arquivos .txt foram lidos com sucesso.{RESET}")
        return all_contents
    else:
        print(f"{DARK_RED}[DEBUG] Nenhum arquivo .txt encontrado.{RESET}")
        return f"{DARK_RED}Nenhum arquivo .txt encontrado em nenhum dos diretórios.{RESET}"
def selecionar_declividade():
    arquivo = 'integracao.xlsx'
    aba_nome = 'DECLIVIDADE E PEDOLOGIA'  # Nome da aba onde os dados estão localizados
    coluna = 2  # Número da coluna que você quer ler (por exemplo, 2 para coluna B)

    # Função para ler a coluna específica da aba
    def ler_coluna_excel(arquivo, aba_nome, coluna):
        try:
            workbook = load_workbook(arquivo)
            aba = workbook[aba_nome]
            valores = []
            for linha in aba.iter_rows(min_col=coluna, max_col=coluna + 1, min_row=3,max_row=8, values_only=True):
                if linha[0] is not None:
                    valores.append((linha[0], linha[1]))  # Adiciona o valor da coluna e a célula da direita
            return valores
        except FileNotFoundError:
            messagebox.showerror("Erro", f"Arquivo '{arquivo}' não encontrado.")
        except KeyError:
            messagebox.showerror("Erro", f"Aba '{aba_nome}' não encontrada no arquivo Excel.")
        return []

    # Função para exibir a lista e permitir que o usuário selecione um ou mais valores
    def selecionar_valor(valores):
        selecoes = []

        def confirmar_selecao():
            selecao = listbox.curselection()
            if selecao:
                for index in selecao:
                    selecionado = valores[index]
                    selecoes.append(selecionado)
                janela.destroy()
            else:
                messagebox.showwarning("Atenção", "Nenhum valor selecionado!")

        # Criação da interface gráfica
        janela = tk.Tk()
        janela.title("Seleção de Valores do Excel")
        janela.geometry("500x400")
        janela.configure(bg="#e6f7ff")

        # Rótulo de instrução
        rotulo_instrucao = ttk.Label(janela, text="Selecione as declividades abaixo:", background="#e6f7ff", font=("Helvetica", 12, "bold"))
        rotulo_instrucao.pack(pady=(10, 5))

        # Listbox com barra de rolagem
        frame_listbox = tk.Frame(janela, bg="#e6f7ff")
        frame_listbox.pack(padx=20, pady=10, fill=tk.BOTH, expand=True)

        scrollbar = tk.Scrollbar(frame_listbox, orient=tk.VERTICAL)
        listbox = tk.Listbox(frame_listbox, selectmode=tk.MULTIPLE, yscrollcommand=scrollbar.set, height=10, width=50, bg="#ffffff", fg="#003366", font=("Arial", 10))
        scrollbar.config(command=listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for valor in valores:
            listbox.insert(tk.END, valor[0])

        # Botão para confirmar a seleção
        botao_confirmar = ttk.Button(janela, text="Confirmar Seleção", command=confirmar_selecao)
        botao_confirmar.pack(pady=(5, 15))

        # Tornar a janela responsiva
        janela.grid_rowconfigure(0, weight=1)
        janela.grid_columnconfigure(0, weight=1)
        frame_listbox.grid_rowconfigure(0, weight=1)
        frame_listbox.grid_columnconfigure(0, weight=1)

        janela.mainloop()
        return selecoes

    # Lê a coluna do Excel e exibe a interface para seleção
    valores_coluna = ler_coluna_excel(arquivo, aba_nome, coluna)
    if valores_coluna:
        return selecionar_valor(valores_coluna)
    else:
        return []
def selecionar_pedologia():
    arquivo = 'integracao.xlsx'
    aba_nome = 'DECLIVIDADE E PEDOLOGIA'  # Nome da aba onde os dados estão localizados
    coluna = 2  # Número da coluna que você quer ler (por exemplo, 2 para coluna B)

    # Função para ler a coluna específica da aba
    def ler_coluna_excel(arquivo, aba_nome, coluna):
        try:
            workbook = load_workbook(arquivo)
            aba = workbook[aba_nome]
            valores = []
            for linha in aba.iter_rows(min_col=coluna, max_col=coluna + 1, min_row=14,max_row=33, values_only=True):
                if linha[0] is not None:
                    valores.append((linha[0], linha[1]))  # Adiciona o valor da coluna e a célula da direita
            return valores
        except FileNotFoundError:
            messagebox.showerror("Erro", f"Arquivo '{arquivo}' não encontrado.")
        except KeyError:
            messagebox.showerror("Erro", f"Aba '{aba_nome}' não encontrada no arquivo Excel.")
        return []

    # Função para exibir a lista e permitir que o usuário selecione um ou mais valores
    def selecionar_valor(valores):
        selecoes = []

        def confirmar_selecao():
            selecao = listbox.curselection()
            if selecao:
                for index in selecao:
                    selecionado = valores[index]
                    selecoes.append(selecionado)
                janela.destroy()
            else:
                messagebox.showwarning("Atenção", "Nenhum valor selecionado!")

        # Criação da interface gráfica
        janela = tk.Tk()
        janela.title("Seleção de Valores do Excel")
        janela.geometry("500x400")
        janela.configure(bg="#e6f7ff")

        # Rótulo de instrução
        rotulo_instrucao = ttk.Label(janela, text="Selecione as pedologias abaixo:", background="#e6f7ff", font=("Helvetica", 12, "bold"))
        rotulo_instrucao.pack(pady=(10, 5))

        # Listbox com barra de rolagem
        frame_listbox = tk.Frame(janela, bg="#e6f7ff")
        frame_listbox.pack(padx=20, pady=10, fill=tk.BOTH, expand=True)

        scrollbar = tk.Scrollbar(frame_listbox, orient=tk.VERTICAL)
        listbox = tk.Listbox(frame_listbox, selectmode=tk.MULTIPLE, yscrollcommand=scrollbar.set, height=10, width=50, bg="#ffffff", fg="#003366", font=("Arial", 10))
        scrollbar.config(command=listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for valor in valores:
            listbox.insert(tk.END, valor[0])

        # Botão para confirmar a seleção
        botao_confirmar = ttk.Button(janela, text="Confirmar Seleção", command=confirmar_selecao)
        botao_confirmar.pack(pady=(5, 15))

        # Tornar a janela responsiva
        janela.grid_rowconfigure(0, weight=1)
        janela.grid_columnconfigure(0, weight=1)
        frame_listbox.grid_rowconfigure(0, weight=1)
        frame_listbox.grid_columnconfigure(0, weight=1)

        janela.mainloop()
        return selecoes

    # Lê a coluna do Excel e exibe a interface para seleção
    valores_coluna = ler_coluna_excel(arquivo, aba_nome, coluna)
    if valores_coluna:
        return selecionar_valor(valores_coluna)
    else:
        return []
def gerar_texto(lista_tuplas,):
    
    texto = ""
    for item in lista_tuplas:
        paragrafo = f"\t<tag>{item[0]}:</tag> {item[1]}\n"
        texto += paragrafo

    
    return texto
def substituir_cabecalho(texto,entrada,saida):
    # Carregar o documento Word
    doc = Document(entrada)

    # Acessar o cabeçalho da seção 1
    section = doc.sections[0]
    header = section.header

    # Modificar apenas o texto do cabeçalho, preservando imagens e formatação
    for paragraph in header.paragraphs:
        if paragraph.text.strip():
            paragraph.clear()
            run = paragraph.add_run(texto)
            run.font.size = Pt(11)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            break

    # Salvar o documento modificado
    doc.save(saida)

    print(f"{DARK_GREEN}O cabeçalho da Seção 1 foi modificado com sucesso.{RESET}")
def extrair_nomes_pedologias(pedologias):
    return [nome.split(' (')[0] for nome, descricao in pedologias]
def fazer_texto_pedologia(lista_pedologias):
    
    if not lista_pedologias:
        raise ValueError(f"{DARK_RED}A lista de pedologias está vazia.{RESET}")

    if len(lista_pedologias) == 1:
        pedologia = lista_pedologias[0]
        texto_atualizado = f"o solo predominante no imóvel é o <tag>{pedologia}</tag>"
    else:
        pedologias_formatadas = " e ".join(", ".join(lista_pedologias).rsplit(", ", 1))
        texto_atualizado = f"os solos predominantes no imóvel são os <tag>{pedologias_formatadas}</tag>"

    return texto_atualizado
def extrair_iniciais_desclividades(data):

    initials = []
    for item in data:
        if isinstance(item, tuple) and len(item) > 0:
            first_element = item[0]
            if isinstance(first_element, str) and len(first_element) > 0:
                initials.append(first_element[0])
    return initials
def fazer_titulo_Declividade(lst):
    # Verifica se a lista é vazia
    if not lst:
        return ""

    # Divide a lista em grupos de 2
    groups = [lst[i:i + 2] for i in range(0, len(lst), 2)]

    # Formata os grupos
    formatted_groups = []
    for group in groups:
        formatted_groups.append("".join(group))

    # Junta os grupos com o "e" ao final
    if len(formatted_groups) > 1:
        return ", ".join(formatted_groups[:-1]) + " e " + formatted_groups[-1]
    else:
        return formatted_groups[0]
def contar_paginas_docx():
    word = 'teste.docx'
    nome_arquivo = word
    # Verifica se o arquivo existe no diretório atual
    if not os.path.isfile(nome_arquivo):
        print(f"{DARK_RED}O arquivo '{nome_arquivo}' não foi encontrado no diretório atual.{RESET}")
        return

    try:
        # Inicializa o Word sem interface visível e sem a interface de COM (que é mais rápida)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # Torna o Word invisível
        word.DisplayAlerts = 0  # Desabilita alertas que podem desacelerar o processo
        
        # Abre o documento sem carregar outros elementos (como fontes, imagens, etc.)
        documento = word.Documents.Open(os.path.abspath(nome_arquivo), ReadOnly=True)

        # Obtém o número de páginas
        numero_paginas = documento.ComputeStatistics(2)  # 2 significa wdStatisticPages
        
        # Fecha o documento e o Word
        documento.Close(False)
        word.Quit()

        # Exibe o número de páginas
        return int(numero_paginas)
    except Exception as e:
        print(f"{DARK_RED}Ocorreu um erro: {e}{RESET}")
        return None
def numero_por_extenso(numero):
    unidades = [
        "zero", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove",
        "dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"
    ]
    dezenas = [
        "", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"
    ]
    centenas = [
        "", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"
    ]

    if int(numero) < 20:
        return unidades[int(numero)]
    elif int(numero) < 100:
        dezena = int(numero) // 10
        unidade = int(numero) % 10
        if unidade == 0:
            return dezenas[dezena]
        else:
            return f"{dezenas[dezena]} e {unidades[unidade]}"
    elif int(numero) < 1000:
        centena = int(numero) // 100
        resto = int(numero) % 100
        if resto == 0:
            return centenas[centena]
        else:
            return f"{centenas[centena]} e {numero_por_extenso(resto)}"
    else:
        return f"{DARK_RED}Número fora do intervalo suportado.{RESET}"    
def colocar_quantidade_de_paginas_laudo_e_enumeradas(quantidade_imagens):
    numero_paginas = contar_paginas_docx()#aqui é pra eu saber quantas paginas tem no laudo
    #aqui nessa função eu to tratando quantas paginas tem no laudo e quantas paginas tem enumeradas no laudo
    numero_outros_Textos = 5
    numero_de_paginas_numeradas = numero_paginas - quantidade_imagens - numero_outros_Textos
    dados_quantidade_pagina = {
        'hd190':texto_paginas_numeradas(numero_paginas),#saida : inteiro (extenso)
        '0jasa':texto_paginas_numeradas(numero_de_paginas_numeradas)#saida : inteiro (extenso)
    }
    entrada_e_saida = 'teste.docx'
    substituir_palavras_documento(entrada_e_saida,dados_quantidade_pagina,entrada_e_saida)
def copiar_ou_recortar_arquivos(origem, destino, acao="copiar"):

    # Verificar se as pastas de origem e destino existem
    if not os.path.exists(origem):
        print(f"{DARK_RED}A pasta de origem '{origem}' não existe.{RESET}")
        return
    if not os.path.exists(destino):
        os.makedirs(destino)

    # Extensões de arquivos a serem manipulados
    extensoes = ['.xlsx', '.xls', '.docx', '.doc']

    # Iterar pelos arquivos na pasta de origem
    for arquivo in os.listdir(origem):
        caminho_completo_origem = os.path.join(origem, arquivo)

        # Verificar se é um arquivo e se tem a extensão desejada
        if os.path.isfile(caminho_completo_origem) and any(arquivo.endswith(ext) for ext in extensoes):
            caminho_completo_destino = os.path.join(destino, arquivo)

            if acao == "copiar":
                # Copiar o arquivo
                shutil.copy2(caminho_completo_origem, caminho_completo_destino)
                print(f"{DARK_GREEN}Arquivo '{arquivo}' copiado para '{destino}'.{RESET}")
            elif acao == "recortar":
                # Mover o arquivo
                shutil.move(caminho_completo_origem, caminho_completo_destino)
                print(f"{DARK_GREEN}Arquivo '{arquivo}' movido para '{destino}'.{RESET}")
            else:
                print(f"{DARK_GREEN}Ação desconhecida: '{acao}'. Use 'copiar' ou 'recortar'.{RESET}")
def encontrar_nomes(lista, nomes):
 
    resultados = {}
    for nome in nomes:
        resultado = next((item for item in lista if nome.lower() in str(item).lower()), None)
        resultados[nome] = resultado
    return resultados
def renomear_arquivo_word(caminho_arquivo, novo_nome):
    # Verificar se o arquivo existe
    if not os.path.exists(caminho_arquivo):
        raise FileNotFoundError(f"{DARK_RED}O arquivo '{caminho_arquivo}' não foi encontrado.{RESET}")
    
    # Verificar se o arquivo tem a extensão .docx
    if not caminho_arquivo.lower().endswith('.docx'):
        raise ValueError(f"{DARK_RED}O arquivo especificado não é um arquivo Word (.docx).{RESET}")
    
    # Obter o diretório do arquivo original
    diretorio = os.path.dirname(caminho_arquivo)
    
    # Garantir que o novo nome também tenha a extensão .docx
    if not novo_nome.lower().endswith('.docx'):
        novo_nome += '.docx'
    
    # Criar o caminho completo para o novo arquivo
    novo_caminho = os.path.join(diretorio, novo_nome)
    
    # Renomear o arquivo
    os.rename(caminho_arquivo, novo_caminho)
    
    return novo_caminho
def negritar_texto_entre_tags(input_file):
    
    # Carrega o documento
    doc = Document(input_file)
    
    # Regex para capturar o que está entre <tag> e </tag>
    pattern = re.compile(r'(<tag>.*?</tag>)')
    
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        
        # Se não existir <tag>...</tag> no parágrafo, pula
        if '<tag>' not in original_text:
            continue
        
        # Divide o texto do parágrafo em pedaços: 
        #  - pedaços fora das tags
        #  - pedaços que correspondem ao padrão (<tag>...</tag>)
        splitted = re.split(pattern, original_text)
        
        # Limpa o texto do parágrafo para construí-lo do zero
        paragraph.text = ''
        
        for parte in splitted:
            # Verifica se a parte corresponde ao padrão <tag>...</tag>
            if pattern.match(parte):
                # Extrai só o conteúdo entre <tag> e </tag>, sem as próprias tags
                conteudo_sem_tags = re.sub(r'</?tag>', '', parte)
                
                # Cria uma Run em negrito
                run = paragraph.add_run(conteudo_sem_tags)
                run.bold = True
            else:
                # Conteúdo que está fora das tags permanece normal
                paragraph.add_run(parte)
                
    # Salva o resultado
    doc.save(input_file)
def enontrar_estatisticas(base_path):
  
    def search_in_directory(directory):
        
        for root, dirs, files in os.walk(directory):
            print(f"{DARK_GREEN}[DEBUG] Verificando o diretório: {root}{RESET}")
            print(f"{DARK_GREEN}[DEBUG] Subdiretórios encontrados: {dirs}{RESET}")
            for dir_name in dirs:
                print(f"{DARK_GREEN}[DEBUG] Verificando subdiretório: {dir_name}{RESET}")
                if 'peças' in dir_name.lower():
                    # Caminho completo da pasta que contém "peças"
                    peças_dir = os.path.join(root, dir_name)
                    print(f"{DARK_GREEN}[DEBUG] Encontrada a pasta 'peças': {peças_dir}{RESET}")

                    # Procura por arquivos Excel na pasta
                    excel_files = []
                    for file_name in os.listdir(peças_dir):
                        print(f"{DARK_GREEN}[DEBUG] Verificando o arquivo: {file_name}{RESET}")
                        if file_name.endswith(('.xls', '.xlsx')):
                            file_path = os.path.join(peças_dir, file_name)
                            print(f"{DARK_GREEN}[DEBUG] Arquivo Excel encontrado: {file_path}{RESET}")
                            excel_files.append(file_path)

                    return excel_files
        return []

    # Caminho inicial
    current_dir = os.path.dirname(os.path.dirname(base_path))

    # Procura no diretório atual e vai subindo até a raiz do sistema
    while current_dir:
        print(f"{DARK_GREEN}[DEBUG] Procurando no diretório: {current_dir}{RESET}")
        excel_files = search_in_directory(current_dir)
        if excel_files:
            return excel_files

        # Sobe um nível na hierarquia
        parent_dir = os.path.dirname(current_dir)
        if parent_dir == current_dir:  # Verifica se atingiu a raiz do sistema
            break
        current_dir = parent_dir

    print(f"{DARK_RED}[DEBUG] Nenhuma pasta 'peças' ou arquivo Excel encontrado.{RESET}")
    return []
def extrair_numeros(texto: str) -> str:
    print(f"{DARK_GREEN}Texto da matricula: {texto}{RESET}")  # Mostra o texto de entrada
    numeros = [c for c in texto if c.isdigit()]  # Lista com os dígitos extraídos
    print(f"{DARK_GREEN}numeros encontrados: {numeros}{RESET}")  # Mostra os dígitos extraídos
    resultado = ''.join(numeros)  # Concatena os dígitos em uma string
    print(f"{DARK_GREEN}Resultado final: {resultado}{RESET}")  # Mostra o resultado final
    return resultado
def gerar_texto_coeficientes(lista1, lista2):
    print(f"{DARK_GREEN}Lista1: {lista1}{RESET}")
    print(f"{DARK_GREEN}Lista2: {lista2}{RESET}")

    if len(lista1) != len(lista2):
        raise ValueError(f"{DARK_GREEN}As duas listas devem ter o mesmo tamanho.{RESET}")

    resultado = []
    for valor1, valor2 in zip(lista1, lista2):
        print(f"{DARK_GREEN}Processando: {valor1}% (Mat. {valor2}){RESET}")
        resultado.append(f"{valor1}% (Mat. {valor2})")

    if len(resultado) > 1:
        texto_final = ", ".join(resultado[:-1]) + " é " + resultado[-1]
    else:
        texto_final = resultado[0]

    print(f"{DARK_GREEN}Texto final formatado: {texto_final}{RESET}")
    return texto_final
def buscar_valor_excel(caminho_arquivo,nome_planilha,endereco_celula):
   
    print(f"{DARK_GREEN}[INFO] Iniciando a função obter_valor_celula...{RESET}")
    print(f"{DARK_GREEN}[INFO] Abrindo o arquivo: {caminho_arquivo}{RESET}")

    # Se quiser ler o TEXTO da FÓRMULA, use data_only=False
    # Se quiser ler apenas o RESULTADO (valor calculado), use data_only=True
    workbook = load_workbook(caminho_arquivo, data_only=True)
    print(f"{DARK_GREEN}[INFO] Arquivo aberto com sucesso.{RESET}")

    print(f"{DARK_GREEN}[INFO] Selecionando a planilha: {nome_planilha}{RESET}")
    planilha = workbook[nome_planilha]
    print(f"{DARK_GREEN}[INFO] Planilha '{nome_planilha}' selecionada.{RESET}")

    print(f"{DARK_GREEN}[INFO] Lendo o valor (ou fórmula) da célula: {endereco_celula}{RESET}")
    valor = planilha[endereco_celula].value

    print(f"{DARK_GREEN}[INFO] Valor encontrado na célula {endereco_celula}: {valor}{RESET}")

    print(f"{DARK_GREEN}[INFO] Finalizando a função obter_valor_celula...\n{RESET}")
    return valor
def colocar_tag_texto(texto):
    return f'<tag>{texto}</tag>'
def caminho_pasta_pecas_tecnicas(excel_path):
  
    print(f"{DARK_GREEN}[DEBUG] Verificando se o arquivo existe: {excel_path}{RESET}")
    if os.path.isfile(excel_path):
        folder_path = os.path.dirname(excel_path)
        print(f"{DARK_GREEN}[DEBUG] Caminho da pasta encontrado: {folder_path}{RESET}")
        return folder_path
    else:
        print(f"{DARK_RED}[ERROR] O arquivo '{excel_path}' não foi encontrado.{RESET}")
        raise FileNotFoundError(f"{DARK_RED}O arquivo '{excel_path}' não foi encontrado.{RESET}")
def procurar_arquivo_word(directory):
    print(f"{DARK_GREEN}Procurando arquivos Word no diretório: {directory}{RESET}")

    # Busca apenas na pasta especificada, sem entrar em subpastas
    try:
        arquivos = os.listdir(directory)
        print(f"{DARK_GREEN}Arquivos encontrados no diretório: {arquivos}{RESET}")
        for file in arquivos:
            print(f"{DARK_GREEN}Verificando arquivo: {file}{RESET}")
            if file.endswith('.docx') or file.endswith('.doc'):
                caminho_completo = os.path.join(directory, file)
                print(f"{DARK_GREEN}Arquivo Word encontrado: {caminho_completo}{RESET}")
                return caminho_completo
        print(f"{DARK_RED}Nenhum arquivo Word encontrado no diretório especificado.{RESET}")
    except Exception as e:
        print(f"{DARK_RED}Erro ao acessar o diretório: {e}{RESET}")

    return None
def texto_paginas_numeradas(numero):
    """
    Recebe um número inteiro e retorna o número completo por extenso entre parênteses.
    
    Exemplo:
    numero_por_extenso(123) => "123 (cento e vinte e três)"
    """
    unidades = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
    dezenas = ["", "dez", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
    centenas = ["", "cem", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
    especiais = {
        11: "onze", 12: "doze", 13: "treze", 14: "catorze", 15: "quinze", 16: "dezesseis", 17: "dezessete", 18: "dezoito", 19: "dezenove"
    }

    def extenso(numero):
        if numero == 0:
            return "zero"

        partes = []

        if numero >= 100:
            centena = numero // 100
            partes.append(centenas[centena] if numero % 100 != 0 or centena != 1 else "cem")
            numero %= 100

        if 10 < numero < 20:
            partes.append(especiais[numero])
        else:
            if numero >= 10:
                dezena = numero // 10
                partes.append(dezenas[dezena])
                numero %= 10

            if numero > 0:
                partes.append(unidades[numero])

        return " e ".join(partes)

    resultado_extenso = extenso(abs(numero))
    if numero < 0:
        resultado_extenso = "menos " + resultado_extenso

    return f"{numero} ({resultado_extenso})"
def encontrar_croqui_no_diretorio(excel_path):
    # 1. Obtém o diretório onde está o arquivo Excel
    dir_excel = os.path.dirname(excel_path)
    
    # 2. Sobe um diretório a partir do diretório do Excel
    dir_pai = os.path.dirname(dir_excel)
    
    # 3. Compila um padrão regex para 'croqui' (ignorando maiúscula/minúscula)
    padrao_croqui = re.compile(r'croqui', re.IGNORECASE)
    
    # 4. Lista para armazenar os caminhos das imagens encontradas
    imagens_encontradas = []

    # 5. Percorre o diretório pai (e seus subdiretórios) em busca de arquivos
    for raiz, _, arquivos in os.walk(dir_pai):
        for arquivo in arquivos:
            # Verifica se o nome do arquivo corresponde ao padrão 'croqui' e tem extensão .jpg ou .png
            if padrao_croqui.search(arquivo) and (arquivo.lower().endswith('.jpg') or arquivo.lower().endswith('.png')):
                print(f'{DARK_GREEN}Imagem do croqui encontrada: {arquivo}{RESET}')
                imagens_encontradas.append(os.path.join(raiz, arquivo))

    # Verifica se encontrou alguma imagem
    if imagens_encontradas:
        print(f'{DARK_GREEN}{len(imagens_encontradas)} imagens de croqui encontradas.{RESET}')
        return imagens_encontradas
    else:
        print(f'{DARK_RED}Nenhuma imagem do croqui encontrada.{RESET}')
        return []
def listar_imagens_na_pasta_mapas(caminho_arquivo_excel):
    # Obtém o diretório do arquivo Excel
    diretorio_excel = os.path.dirname(caminho_arquivo_excel)

    # Vai um diretório acima
    diretorio_superior = os.path.abspath(os.path.join(diretorio_excel, os.pardir))

    # Caminho da pasta "MAPAS" no diretório superior
    pasta_mapas = os.path.join(diretorio_superior, "MAPAS")

    # Verifica se a pasta existe
    if not os.path.exists(pasta_mapas):
        raise FileNotFoundError(f"A pasta 'MAPAS' não foi encontrada no diretório {diretorio_superior}.")

    # Lista todos os arquivos na pasta "MAPAS"
    arquivos = os.listdir(pasta_mapas)

    # Filtra apenas imagens PNG e JPG
    imagens = [
        os.path.join(pasta_mapas, arquivo)
        for arquivo in arquivos
        if arquivo.lower().endswith((".png", ".jpg"))
    ]

    return imagens
def enontrar_imagens_documentos(excel_file_path, folder_name):
    
    # Obtém o diretório base (um nível acima do arquivo Excel)
    base_directory = os.path.abspath(os.path.join(os.path.dirname(excel_file_path), '..'))
    print(f"{DARK_GREEN}Diretório base calculado: {base_directory}{RESET}")

    # Caminho para a pasta DOCUMENTOS/PNG
    png_directory = os.path.join(base_directory, "DOCUMENTOS", "PNG")
    print(f"{DARK_GREEN}Caminho para a pasta PNG: {png_directory}{RESET}")

    if not os.path.exists(png_directory):
        print(f"{DARK_RED}A pasta 'PNG' não foi encontrada em {png_directory}{RESET}")
        return None

    # Lista todas as subpastas dentro da pasta PNG
    subfolders = [f for f in os.listdir(png_directory) if os.path.isdir(os.path.join(png_directory, f))]
    print(f"{DARK_GREEN}Subpastas encontradas: {subfolders}{RESET}")

    # Filtra as subpastas que contêm o nome fornecido (case insensitive)
    matching_folders = [sf for sf in subfolders if folder_name.lower() in sf.lower()]
    print(f"{DARK_GREEN}Pastas que correspondem ao filtro '{folder_name}': {matching_folders}{RESET}")

    if not matching_folders:
        print(f"{DARK_RED}Nenhuma pasta contendo {folder_name} foi encontrada.{RESET}")
        return None

    all_image_files = []

    # Itera sobre todas as pastas correspondentes
    for target_folder in matching_folders:
        target_folder_path = os.path.join(png_directory, target_folder)
        print(f"{DARK_GREEN}Caminho completo da pasta alvo: {target_folder_path}{RESET}")

        # Lista todas as imagens na pasta atual
        image_files = [
            os.path.join(target_folder_path, f)
            for f in os.listdir(target_folder_path)
            if os.path.isfile(os.path.join(target_folder_path, f)) and f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif'))
        ]
        print(f"{DARK_GREEN}as Imagens foram encontradas na pasta {target_folder}{RESET}")

        all_image_files.extend(image_files)

    return all_image_files if all_image_files else None
def encontrar_primeiro_excel_pecas(lista_diretorios, nome_diretorio):
  
    # Configuração do logging
    # Ajuste o level para DEBUG ou INFO conforme o detalhamento desejado.
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%H:%M:%S'
    )

    # 1) Itera sobre cada diretório base
    for caminho_base in lista_diretorios:
        logging.info(f"{DARK_GREEN}Analisando o diretório base: {caminho_base}{RESET}")

        # Verifica se o caminho_base existe e é diretório
        if not os.path.isdir(caminho_base):
            logging.warning(f"{DARK_RED}Não é um diretório válido ou não existe: {caminho_base}{RESET}")
            continue

        try:
            # 2) Listamos o primeiro nível dentro do caminho_base
            itens_primeiro_nivel = os.listdir(caminho_base)
        except PermissionError:
            logging.warning(f"{DARK_RED}Sem permissão para acessar: {caminho_base}{RESET}")
            continue
        except FileNotFoundError:
            logging.warning(f"{DARK_RED}Diretório não encontrado: {caminho_base}{RESET}")
            continue

        # 3) Para cada item desse primeiro nível, checamos se é uma pasta que contém o nome do processo
        for item in itens_primeiro_nivel:
            caminho_item = os.path.join(caminho_base, item)
            logging.debug(f"{DARK_GREEN}Verificando item no primeiro nível: {caminho_item}{RESET}")

            # Verifica se é diretório e se o nome_diretorio está no nome dele
            if os.path.isdir(caminho_item) and (nome_diretorio in item):
                logging.info(f"{DARK_GREEN}Encontrou subdiretório com '{nome_diretorio}': {caminho_item}{DARK_GREEN}")

                # 4) Dentro desta pasta, listamos o primeiro nível de novo
                try:
                    itens_subpasta = os.listdir(caminho_item)
                except PermissionError:
                    logging.warning(f"{DARK_RED}Sem permissão para acessar: {caminho_item}{RESET}")
                    continue
                except FileNotFoundError:
                    logging.warning(f"{DARK_RED}Subdiretório não encontrado: {caminho_item}{DARK_GREEN}")
                    continue

                # 5) Procuramos a pasta "PEÇAS"
                for subitem in itens_subpasta:
                    caminho_subitem = os.path.join(caminho_item, subitem)
                    logging.debug(f"{DARK_GREEN}Verificando subitem dentro de {caminho_item}: {caminho_subitem}{RESET}")

                    # Se for diretório e contiver "PEÇAS" no nome
                    if os.path.isdir(caminho_subitem) and ("PEÇAS" in subitem.upper()):
                        logging.info(f"{DARK_GREEN}Encontrou pasta 'PEÇAS': {caminho_subitem}{RESET}")

                        # 6) Listar arquivos dentro da pasta "PEÇAS"
                        try:
                            arquivos_pecas = os.listdir(caminho_subitem)
                        except PermissionError:
                            logging.warning(f"{DARK_RED}Sem permissão para acessar: {caminho_subitem}{RESET}")
                            continue
                        except FileNotFoundError:
                            logging.warning(f"{DARK_RED}Subdiretório não encontrado: {caminho_subitem}{RESET}")
                            continue

                        # 7) Verificar o primeiro arquivo Excel
                        for arquivo in arquivos_pecas:
                            logging.debug(f"{DARK_GREEN}Analisando arquivo: {arquivo}{RESET}")
                            if arquivo.lower().endswith(('.xlsx', '.xls', '.xlsm')):
                                caminho_excel = os.path.join(caminho_subitem, arquivo)
                                logging.info(f"{DARK_GREEN}Encontrou arquivo Excel: {caminho_excel}{RESET}")
                                return caminho_excel

                # Se chegou aqui, não encontrou a pasta "PEÇAS" ou nenhum Excel
                # Mas continua a procurar em outros subdiretórios que tenham nome_diretorio.
                logging.debug(f"{DARK_RED}Nenhuma pasta 'PEÇAS' (com Excel) encontrada em: {caminho_item}{RESET}")

    # Se passarmos por todos os diretórios e não encontrarmos nada
    logging.info(f"{DARK_RED}Nenhum arquivo Excel encontrado com '{nome_diretorio}' na lista de diretórios.{RESET}")
    return None
def escolha_usuario():
    def user_choice():
        # Função interna para capturar a escolha do usuário
        choice = var.get()
        if choice:

            nonlocal resultado
            resultado = choice  # Atualiza o resultado com a escolha do usuário
            root.destroy()  # Fecha a interface após a escolha
        else:
            messagebox.showwarning("Atenção", "Por favor, escolha uma opção!")

    # Variável para armazenar o resultado
    resultado = None

    # Criar a janela principal
    root = tk.Tk()
    root.title("Escolha uma Opção")

    # Variável para armazenar a escolha do usuário
    var = tk.StringVar(value="")

    # Criar os botões de opção
    tk.Label(root, text="Escolha uma das opções:").pack(pady=10)
    tk.Radiobutton(root, text="Declividade", variable=var, value="declividade").pack(anchor="w")
    tk.Radiobutton(root, text="vegetação", variable=var, value="vegetação").pack(anchor="w")

    # Botão para confirmar a escolha
    tk.Button(root, text="Confirmar", command=user_choice).pack(pady=10)

    # Iniciar o loop da interface
    root.mainloop()

    # Retornar o resultado escolhido
    return resultado
def selecionar_vegetacao():
    arquivo = 'integracao.xlsx'
    aba_nome = 'DECLIVIDADE E PEDOLOGIA'  # Nome da aba onde os dados estão localizados
    coluna = 2  # Número da coluna que você quer ler (por exemplo, 2 para coluna B)

    # Função para ler a coluna específica da aba
    def ler_coluna_excel(arquivo, aba_nome, coluna):
        try:
            workbook = load_workbook(arquivo)
            aba = workbook[aba_nome]
            valores = []
            for linha in aba.iter_rows(min_col=coluna, max_col=coluna + 1, min_row=39,max_row=57, values_only=True):
                if linha[0] is not None:
                    valores.append((linha[0], linha[1]))  # Adiciona o valor da coluna e a célula da direita
            return valores
        except FileNotFoundError:
            messagebox.showerror("Erro", f"Arquivo '{arquivo}' não encontrado.")
        except KeyError:
            messagebox.showerror("Erro", f"Aba '{aba_nome}' não encontrada no arquivo Excel.")
        return []

    # Função para exibir a lista e permitir que o usuário selecione um ou mais valores
    def selecionar_valor(valores):
        selecoes = []

        def confirmar_selecao():
            selecao = listbox.curselection()
            if selecao:
                for index in selecao:
                    selecionado = valores[index]
                    selecoes.append(selecionado)
                janela.destroy()
            else:
                messagebox.showwarning("Atenção", "Nenhum valor selecionado!")

        # Criação da interface gráfica
        janela = tk.Tk()
        janela.title("Seleção de Valores do Excel")
        janela.geometry("500x400")
        janela.configure(bg="#e6f7ff")

        # Rótulo de instrução
        rotulo_instrucao = ttk.Label(janela, text="Selecione as pedologias abaixo:", background="#e6f7ff", font=("Helvetica", 12, "bold"))
        rotulo_instrucao.pack(pady=(10, 5))

        # Listbox com barra de rolagem
        frame_listbox = tk.Frame(janela, bg="#e6f7ff")
        frame_listbox.pack(padx=20, pady=10, fill=tk.BOTH, expand=True)

        scrollbar = tk.Scrollbar(frame_listbox, orient=tk.VERTICAL)
        listbox = tk.Listbox(frame_listbox, selectmode=tk.MULTIPLE, yscrollcommand=scrollbar.set, height=10, width=50, bg="#ffffff", fg="#003366", font=("Arial", 10))
        scrollbar.config(command=listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        for valor in valores:
            listbox.insert(tk.END, valor[0])

        # Botão para confirmar a seleção
        botao_confirmar = ttk.Button(janela, text="Confirmar Seleção", command=confirmar_selecao)
        botao_confirmar.pack(pady=(5, 15))

        # Tornar a janela responsiva
        janela.grid_rowconfigure(0, weight=1)
        janela.grid_columnconfigure(0, weight=1)
        frame_listbox.grid_rowconfigure(0, weight=1)
        frame_listbox.grid_columnconfigure(0, weight=1)

        janela.mainloop()
        return selecoes

    # Lê a coluna do Excel e exibe a interface para seleção
    valores_coluna = ler_coluna_excel(arquivo, aba_nome, coluna)
    if valores_coluna:
        return selecionar_valor(valores_coluna)
    else:
        return []
def extrair_nomes_vegetacao(vegetacao):
    return [nome.split(' (')[0] for nome, descricao in vegetacao]
def apagar_pycache():#TODO : trocar pelo pc do user
    caminho = r'C:\\Users\\Usuario\\Desktop\\nova_pasta_descritivo\\'
    if not os.path.isdir(caminho):
        print(f"{DARK_RED}Erro: O caminho fornecido '{caminho}' não é um diretório válido.{RESET}")
        return
    
    for root, dirs, files in os.walk(caminho):
        for dir_name in dirs:
            if dir_name == "__pycache__":
                caminho_pycache = os.path.join(root, dir_name)
                try:
                    shutil.rmtree(caminho_pycache)
                    print(f"{DARK_RED}Pasta apagada: {caminho_pycache}{RESET}")
                except Exception as e:
                    print(f"{DARK_RED}Erro ao apagar {caminho_pycache}: {e}{RESET}")
def extrair_cidade_uf(input_string):
    try:
        print(f"{DARK_GREEN}[LOG] Iniciando a extração de cidade e UF.{RESET}")

        # Remove espaços em branco extras
        input_string = input_string.strip()
        print(f"{DARK_GREEN}[LOG] Entrada processada: '{input_string}'{RESET}")

        # Verifica se o formato está correto
        if '-' not in input_string:
            print(f"{DARK_RED}[LOG] Formato incorreto detectado.{RESET}")
            raise ValueError(f"O formato deve ser 'cidade - UF'.")

        # Divide a string em cidade e UF baseado no último traço
        partes = input_string.rsplit('-', 1)
        cidade = partes[0].strip()
        uf = partes[1].strip()
        print(f"{DARK_GREEN}[LOG] Cidade extraída: '{cidade}', UF extraída: '{uf}'{RESET}")

        # Validações básicas
        if len(uf) != 2:
            print(f"{DARK_GREEN}[LOG] UF com tamanho inválido.{RESET}")
            raise ValueError(f"UF deve conter exatamente 2 caracteres.")

        print("[LOG] Extração concluída com sucesso.")
        return cidade, uf.upper()
    except ValueError as e:
        print(f"[LOG] Erro durante a extração: {e}")
        return str(e)
def escolher_Texto_mosaicos():
    resultado = []

    def selecionar_opcao(descricao, janela):
        resultado.append(descricao)
        janela.destroy()

    janela = tk.Tk()
    janela.title("Menu de Mosaico")

    label = tk.Label(janela, text="CASO NÃO HOUVER MOSAICO \nAPENAS APERTE NO (X)", font=("Arial", 14))
    label.pack(pady=10)

    botoes = [
        ("AB", "AB – Mosaico com predomínio de A sobre B"),
        ("BA", "BA – Mosaico com predomínio de B sobre A"),
        ("BC", "BC – Mosaico com predomínio de B sobre C"),
        ("CB", "CB – Mosaico com predomínio de C sobre B"),
        ("CD", "CD – Mosaico com predomínio de C sobre D"),
        ("DC", "DC – Mosaico com predomínio de D sobre C")
    ]

    for codigo, descricao in botoes:
        botao = tk.Button(janela, text=descricao, font=("Arial", 12), command=lambda d=descricao: selecionar_opcao(d, janela))
        botao.pack(fill=tk.X, pady=5, padx=20)

    janela.mainloop()
    return "<tag>"+resultado[0]+"</tag>" if resultado else None
def substituir_indices_4_e_5(lista_tuplas, lista_valores_4, lista_valores_5):
    # Verificar se as listas de valores têm o mesmo tamanho que a lista de tuplas
    if len(lista_tuplas) != len(lista_valores_4) or len(lista_tuplas) != len(lista_valores_5):
        print("As listas de valores devem ter o mesmo tamanho que a lista de tuplas.")
        return lista_tuplas
    
    # Percorrer cada tupla e substituir os valores nos índices 4 e 5
    for i in range(len(lista_tuplas)):
        tupla = lista_tuplas[i]
        
        # Verificar se a tupla tem ao menos 6 elementos (para garantir que os índices 4 e 5 existem)
        if len(tupla) > 5:
            # Substituir os valores nos índices 4 e 5 pelas listas de valores fornecidas
            nova_tupla = tupla[:4] + (lista_valores_4[i], lista_valores_5[i]) + tupla[6:]
            lista_tuplas[i] = nova_tupla  # Substituir a tupla na lista
    
    return lista_tuplas
def somar_valores(lista):
    
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    # Converter os valores da lista para números
    valores = [
        locale.atof(valor.replace('R$', '').strip()) for valor in lista
    ]

    # Somar os valores
    soma = sum(valores)

    # Retornar a soma formatada como string
    return f"R$ {locale.format_string('%.2f', soma, grouping=True)}"
def separar_ultimo_item_por_e(items):
  
    print(f"{DARK_GREEN}Entrada: {items}{RESET}")  # Log da entrada
    if not items:
        print(f"{DARK_RED}A lista está vazia. Retornando string vazia.{RESET}")  # Log para lista vazia
        return ""
    elif len(items) == 1:
        print(f"{DARK_GREEN}A lista possui apenas um elemento: {items[0]}{RESET}")  # Log para lista com um item
        return items[0]
    else:
        resultado = ", ".join(items[:-1]) + f" e {items[-1]}"
        print(f"{DARK_GREEN}Lista formatada: {resultado}{RESET}")  # Log do resultado formatado
        return resultado
def agrupar_Matricula_por_Car(lista):
    grupos = defaultdict(list)
    indices_8 = set()

    # Pega o valor do índice 8 da primeira tupla
    valor_padrao = lista[0][8] if lista[0][8] is not None else lista[0][0]

    for elemento in lista:
        # Usa o índice 8 ou o valor da primeira tupla se índice 8 for None
        chave = elemento[8] if elemento[8] is not None else valor_padrao
        grupos[chave].append(elemento[0])  # Adiciona o índice 0 ao grupo correspondente
        indices_8.add(chave)  # Adiciona o índice 8 ao conjunto de chaves

    if len(indices_8) == 1:  # Verifica se todos os índices 8 são iguais
        return "todos itens sao iguais"

    texto = []
    for chave, valores in grupos.items():
        texto.append(f"O car de <tag>{chave}</tag> se refere às matrícula(s): {', '.join(map(str, valores))}.")

    return texto
def fazer_Texto_atraves_De_lista(lista):
    
    if not isinstance(lista, list):
        raise ValueError("O argumento deve ser uma lista.")

    return "\n".join(f"{item}\t" for item in lista)
def agrupar_Matricula_por_data_Emissao(lista,cidade_cartorio):
    grupos = defaultdict(list)
    indices_7 = set()

    # Pega o valor do índice 7 da primeira tupla
    valor_padrao = lista[0][7] if lista[0][7] is not None else lista[0][0]

    for elemento in lista:
        # Usa o índice 7 ou o valor da primeira tupla se índice 7 for None
        chave = elemento[7] if elemento[7] is not None else valor_padrao
        grupos[chave].append(elemento[0])  # Adiciona o índice 0 ao grupo correspondente
        indices_7.add(chave)  # Adiciona o índice 7 ao conjunto de chaves

    if len(indices_7) == 1:  # Verifica se todos os índices 7 são iguais
        return "todos itens sao iguais"

    texto = []
    for chave, valores in grupos.items():
        texto.append(f"matrícula(s): <tag>{', '.join(map(str, valores))}</tag> do Cartório de {cidade_cartorio}, emitida no dia <tag>{formatar_data(str(chave))}</tag>.")

    return texto
def ajustar_milhar(valor):
    # Extrai o número da string
    match = re.search(r'R\$\s*([0-9\.]+,[0-9]{2})', valor)
    if not match:
        raise ValueError("Formato inválido de valor monetário")

    numero_str = match.group(1)

    # Substitui pontos por nada e vírgula por ponto para manipulação numérica
    numero_float = float(numero_str.replace('.', '').replace(',', '.'))

    # Obtém o número da centena
    centena = (int(numero_float) // 100) % 10

    if centena >= 5:
        # Incrementa a unidade de milhar e zera as centenas
        numero_float += 1000 - (numero_float % 1000)
    else:
        # Apenas zera as centenas
        numero_float -= numero_float % 1000

    # Formata o número de volta para o formato desejado
    numero_formatado = f"R$ {numero_float:,.2f}".replace('.', 'X').replace(',', '.').replace('X', ',')

    return numero_formatado
