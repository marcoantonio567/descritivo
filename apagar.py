import re
from docx import Document

def negritar_texto_entre_tags(input_file):
    
    # Carrega o documento
    doc = Document(input_file)
    
    # Regex para capturar o que está entre <tag> e </tag>
    pattern = re.compile(r'(<tag>.*?</tag>)')
    
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        
        # Se não existir <tag>...</tag> no parágrafo, pula
        if '<tag>' not in original_text:
            continue
        
        # Divide o texto do parágrafo em pedaços: 
        #  - pedaços fora das tags
        #  - pedaços que correspondem ao padrão (<tag>...</tag>)
        splitted = re.split(pattern, original_text)
        
        # Limpa o texto do parágrafo para construí-lo do zero
        paragraph.text = ''
        
        for parte in splitted:
            # Verifica se a parte corresponde ao padrão <tag>...</tag>
            if pattern.match(parte):
                # Extrai só o conteúdo entre <tag> e </tag>, sem as próprias tags
                conteudo_sem_tags = re.sub(r'</?tag>', '', parte)
                
                # Cria uma Run em negrito
                run = paragraph.add_run(conteudo_sem_tags)
                run.bold = True
            else:
                # Conteúdo que está fora das tags permanece normal
                paragraph.add_run(parte)
                
    # Salva o resultado
    doc.save(input_file)

