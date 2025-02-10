from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageOps
from docx.shared import Inches
from docx.shared import Pt, Cm
from docx import Document
import os
import io

def colar_maps(doc_path, images_with_placeholders, width=Inches(5.91), height=Inches(4.18)):
  
    # Carrega o documento Word
    doc = Document(doc_path)

    for image_path, placeholder in images_with_placeholders:
        # Adiciona uma borda preta à imagem
        image_with_border_path = os.path.join(os.path.dirname(image_path), "bordered_" + os.path.basename(image_path))
        with Image.open(image_path) as img:
            bordered_image = ImageOps.expand(img, border=2, fill="black")
            bordered_image.save(image_with_border_path)

        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Remove o marcador de posição
                paragraph.text = paragraph.text.replace(placeholder, "")

                # Adiciona a imagem abaixo do parágrafo
                run = paragraph.add_run()
                run.add_picture(image_with_border_path, width=width, height=height)

        # Remove a imagem temporária
        if os.path.exists(image_with_border_path):
            os.remove(image_with_border_path)

    # Salva o documento atualizado
    doc.save(doc_path)
def delete_last_page(input_file):
    # Abrir o documento Word
    doc = Document(input_file)

    # Obter todos os parágrafos do documento
    paragraphs = doc.paragraphs

    # Se o documento tiver conteúdo
    if paragraphs:
        # Identificar o índice do último parágrafo
        last_paragraph_index = len(paragraphs) - 1
        # Remover o último parágrafo
        last_paragraph = paragraphs[last_paragraph_index]
        last_paragraph._element.getparent().remove(last_paragraph._element)
        
    # Salvar o documento modificado
    doc.save(input_file)
def adcionar_imagens_documentos(file_path, image_paths, titulo):
    # Load the Word document
    doc = Document(file_path)
    
    def add_border_to_image(image_path, border=2, color='black'):
        """Adds a border to an image and returns the modified image."""
        img = Image.open(image_path)
        if img.mode != "RGB":
            img = img.convert("RGB")  # Converte para RGB se não estiver nesse modo
        img_with_border = ImageOps.expand(img, border=border, fill=color)
        return img_with_border

    # Add a title to the first page
    if image_paths:
        title = doc.add_paragraph()
        title_run = title.add_run(titulo)
        title_run.font.name = "Century Gothic"
        title_run.font.size = Pt(11)
        title_run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the first image on the same page as the title
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        temp_image_path = "temp_image1.png"
        img_with_border = add_border_to_image(image_paths[0])
        img_with_border.save(temp_image_path)
        run.add_picture(temp_image_path, width=Cm(14.5), height=Cm(19.79))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        os.remove(temp_image_path)  # Delete the temporary image

        # Remove the first image from the list
        image_paths = image_paths[1:]

    for index, image_path in enumerate(image_paths):
        # Add a page break
        doc.add_page_break()

        # Add the image, centered on the page
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        temp_image_path = f"temp_image{index + 2}.png"
        img_with_border = add_border_to_image(image_path)
        img_with_border.save(temp_image_path)
        run.add_picture(temp_image_path, width=Cm(14.5), height=Cm(19.79))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        os.remove(temp_image_path)  # Delete the temporary image

    # Add a page break after the last page
    doc.add_page_break()

    # Save the document
    doc.save(file_path)
def substituir_croqui_e_rota_acesso(doc_path, placeholder, images, texts, width_cm, height_cm, border_size=3, border_color="black"):

    def add_border_to_image_in_memory(image_path, border_size, border_color):
        """
        Adiciona uma borda a uma imagem e retorna o conteúdo processado em memória.
        """
        img = Image.open(image_path)
        img_with_border = ImageOps.expand(img, border=border_size, fill=border_color)
        img_bytes = io.BytesIO()
        img_with_border.save(img_bytes, format='JPEG')
        img_bytes.seek(0)  # Reseta o ponteiro do buffer
        return img_bytes

    # Abre o documento existente
    doc = Document(doc_path)

    # Verifica se as listas de imagens e textos têm o mesmo tamanho
    if len(images) != len(texts):
        raise ValueError("As listas de imagens e textos devem ter o mesmo tamanho.")

    # Percorre os parágrafos para encontrar o marcador
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Remove o texto do marcador
            paragraph.text = ""

            for i in range(len(images)):
                # Adiciona o texto em um novo parágrafo antes da imagem
                text_paragraph = paragraph.insert_paragraph_before("\n" + "\t" + texts[i] + "\t" + "\n")
                
                # Alinha o texto à esquerda
                text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Adiciona a imagem em um novo parágrafo
                image_paragraph = paragraph.insert_paragraph_before()
                run = image_paragraph.add_run()

                # Processa a imagem com borda em memória
                bordered_image = add_border_to_image_in_memory(images[i], border_size, border_color)

                # Insere a imagem processada no documento
                run.add_picture(bordered_image, width=Cm(width_cm), height=Cm(height_cm))

                # Centraliza a imagem
                image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            break  # Para após encontrar o marcador

    # Salva o documento atualizado
    doc.save(doc_path)

